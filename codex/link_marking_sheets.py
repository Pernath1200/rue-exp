#!/usr/bin/env python3
"""link_marking_sheets.py — make the Topic Reference column clickable.

Rewrites each mapped Topic Reference cell to the APP UNIT NAME and hyperlinks it
to that unit's deep link. Labels with no unit are left exactly as they are, plain
and unlinked — deliberate, per James 2026-08-14: a link that goes nowhere useful
is worse than none, and the blanks are what say which units to build.

    py -X utf8 codex/link_marking_sheets.py <sheet.docx> [...] [--pdf]

NEVER overwrites a master. Writes `<name> (linked).docx` beside the original.
With --pdf, exports through headless LibreOffice, which preserves the links.

This is label stamping on already-approved content, not authoring — the error,
correction and note columns are not touched.
"""
from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
SITE = "https://pernath1200.github.io/rue-exp/"
SOFFICE = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
EDGE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")

LINK_COLOUR = "0563C1"  # conventional hyperlink blue — must read as clickable


def unit_labels() -> dict[str, tuple[str, str]]:
    """node id -> (student-facing name, url)."""
    out: dict[str, tuple[str, str]] = {}
    for domain in ("grammar", "vocab"):
        for p in (ROOT / "data" / domain / "blocks").glob("*.json"):
            d = json.loads(p.read_text(encoding="utf-8"))
            node = d.get("tree_node") or d.get("id")
            out[node] = (f"{d.get('title')} ({d.get('level')})", f"{SITE}#{node}")
    return out


def set_cell_link(cell, text: str, url: str) -> None:
    """Replace a table cell's content with one hyperlinked run.

    Clearing via `para.runs` is NOT enough — it leaves text held in runs nested
    inside other inline elements, and the first attempt produced cells reading
    "Nouns & Determiners → …Articles (advanced) (B1)". Strip every <w:p> off the
    <w:tc> and build one fresh paragraph instead.
    """
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)
    new_p = OxmlElement("w:p")
    tc.append(new_p)
    para = cell.paragraphs[0]

    rid = para.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), LINK_COLOUR)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(colour)
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    para._p.append(link)


def patch(path: Path, mapping: dict, units: dict) -> tuple[Path, int, int]:
    doc = docx.Document(str(path))
    linked = skipped = 0

    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        idx = next((i for i, h in enumerate(header) if "topic" in h), None)
        if idx is None:
            continue
        for row in table.rows[1:]:
            if len(row.cells) <= idx:
                continue
            cell = row.cells[idx]
            label = cell.text.strip()
            if not label or label in {"—", "-", "–"}:
                continue
            node = mapping.get(label)
            if not node:
                skipped += 1          # unmapped or deliberately unlinked
                continue
            if node not in units:
                skipped += 1
                continue
            name, url = units[node]
            set_cell_link(cell, name, url)
            linked += 1

    out = path.with_name(f"{path.stem} (linked).docx")
    doc.save(str(out))
    return out, linked, skipped


def to_pdf(path: Path) -> Path | None:
    """docx -> html (LibreOffice) -> pdf (Edge). Two hops, on purpose.

    LibreOffice's headless `--convert-to pdf` DROPS every hyperlink annotation.
    Verified 2026-08-14 on a one-line test document with a single link: 0 /URI
    in the output, with and without an explicit writer_pdf_Export filter and a
    clean user profile. It is not the document — the docx this script writes
    carries correct w:hyperlink elements and relationships.

    That silently affected work already done: none of the existing feedback PDFs
    carry clickable links, including the ones whose .docx holds James's Padlet
    deep links.

    LibreOffice's HTML export keeps <a href>, and Edge's print-to-pdf keeps link
    annotations, so the two together do what neither does alone.
    """
    if not SOFFICE.exists():
        print("  !! LibreOffice not found — skipping PDF")
        return None
    subprocess.run(
        [str(SOFFICE), "--headless", "--norestore", "--convert-to", "html",
         "--outdir", str(path.parent), str(path)],
        capture_output=True, text=True, timeout=180,
    )
    html = path.with_suffix(".html")
    if not html.exists():
        print("  !! html step failed")
        return None
    pdf = path.with_suffix(".pdf")
    if not EDGE.exists():
        print("  !! Edge not found — skipping PDF")
        return None
    subprocess.run(
        [str(EDGE), "--headless=new", "--disable-gpu", "--no-pdf-header-footer",
         f"--print-to-pdf={pdf}", html.resolve().as_uri()],
        capture_output=True, text=True, timeout=180,
    )
    html.unlink(missing_ok=True)
    return pdf if pdf.exists() else None


def pdf_link_count(pdf: Path) -> int:
    """Count link annotations, decompressing streams — the whole point is links."""
    import re
    import zlib

    raw = pdf.read_bytes()
    blob = bytearray(raw)
    for m in re.finditer(rb"stream\r?\n", raw):
        s = m.end()
        e = raw.find(b"endstream", s)
        try:
            blob += zlib.decompress(raw[s:e])
        except Exception:  # noqa: BLE001
            pass
    return bytes(blob).count(b"/URI")


def main() -> int:
    args = [a for a in sys.argv[1:] if a != "--pdf"]
    want_pdf = "--pdf" in sys.argv[1:]
    if not args:
        print(__doc__)
        return 2

    mapping = json.loads(
        (ROOT / "codex" / "marking_topic_map.json").read_text(encoding="utf-8")
    )["map"]
    units = unit_labels()

    total_l = total_s = 0
    for a in args:
        p = Path(a)
        if not p.exists():
            print(f"  !! missing {p}")
            continue
        lock = p.with_name(f".~lock.{p.name}#")
        if lock.exists():
            print(f"  !! {p.name} is OPEN in LibreOffice — close it first, skipping")
            continue
        out, linked, skipped = patch(p, mapping, units)
        total_l += linked
        total_s += skipped
        msg = f"  {linked:>2} linked · {skipped:>2} left plain   {out.name}"
        if want_pdf:
            pdf = to_pdf(out)
            if pdf:
                msg += f"   -> pdf, {pdf_link_count(pdf)} live links"
            else:
                msg += "   (pdf failed)"
        print(msg)

    print(f"\n{total_l} cells linked · {total_s} left plain (no unit exists)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
